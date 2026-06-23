from __future__ import annotations

import os
import shutil
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch
import types
import sys
import json


class TestDocumentPushBehavior(unittest.TestCase):
    @patch("src.gateway.wecom_outbound.send_file")
    @patch("src.gateway.wecom_outbound.send_file_card")
    @patch("src.tools.document_tool.generate_report")
    def test_generate_report_handler_attempts_send_file_when_card_succeeds(
        self,
        mock_generate_report: MagicMock,
        mock_send_file_card: MagicMock,
        mock_send_file: MagicMock,
    ) -> None:
        from src.tools.builtin import _generate_report_handler

        mock_generate_report.return_value = "/tmp/report.docx"
        mock_send_file_card.return_value = True

        result = _generate_report_handler(
            {
                "title": "车间通行管理规定",
                "content": "这是一个足够长的文档内容。" * 20,
                "from": "u123",
                "format": "docx",
            }
        )

        self.assertEqual(result, "")
        mock_send_file_card.assert_called_once()
        mock_send_file.assert_called_once_with("u123", "/tmp/report.docx")

    @patch("src.gateway.wecom_outbound.send_file")
    @patch("src.gateway.wecom_outbound.send_file_card")
    @patch("src.tools.document_tool.generate_report")
    def test_generate_report_handler_returns_empty_text_when_push_succeeds(
        self,
        mock_generate_report: MagicMock,
        mock_send_file_card: MagicMock,
        mock_send_file: MagicMock,
    ) -> None:
        from src.tools.builtin import _generate_report_handler

        mock_generate_report.return_value = "/tmp/report.docx"
        mock_send_file_card.return_value = True

        result = _generate_report_handler(
            {
                "title": "车间通行管理规定",
                "content": "这是一个足够长的文档内容。" * 20,
                "from": "u123",
                "format": "docx",
            }
        )

        self.assertEqual(result, "")
        mock_send_file.assert_called_once_with("u123", "/tmp/report.docx")

    @patch("src.gateway.wecom_outbound.send_file")
    @patch("src.gateway.wecom_outbound.send_file_card")
    @patch("src.tools.document_tool.generate_report")
    def test_generate_report_handler_for_wecom_bot_skips_app_push_and_returns_chat_reply(
        self,
        mock_generate_report: MagicMock,
        mock_send_file_card: MagicMock,
        mock_send_file: MagicMock,
    ) -> None:
        from src.tools.builtin import _generate_report_handler

        mock_generate_report.return_value = "/tmp/report.docx"

        result = _generate_report_handler(
            {
                "title": "车间通行管理规定",
                "content": "这是一个足够长的文档内容。" * 20,
                "from": "u123",
                "format": "docx",
                "_source_provider": "wecom_bot",
            }
        )

        self.assertIn("文档已生成", result)
        self.assertIn("http://127.0.0.1:18092/api/v1/documents/report.docx", result)
        mock_send_file_card.assert_not_called()
        mock_send_file.assert_not_called()

    @patch("src.tools.document_tool.generate_report")
    def test_generate_report_handler_for_wecom_bot_returns_file_marker(
        self,
        mock_generate_report: MagicMock,
    ) -> None:
        from src.tools.builtin import _generate_report_handler

        mock_generate_report.return_value = "/tmp/report.docx"

        result = _generate_report_handler(
            {
                "title": "车间通行管理规定",
                "content": "这是一个足够长的文档内容。" * 20,
                "from": "u123",
                "format": "docx",
                "_source_provider": "wecom_bot",
            }
        )

        self.assertIn("[BOT_FILE]", result)
        self.assertIn("/tmp/report.docx", result)

    @patch("src.tools.document_tool.generate_report")
    @patch("src.config.bootstrap.build_settings_service")
    def test_generate_report_handler_falls_back_to_request_text_when_enrichment_times_out(
        self,
        mock_build_settings_service: MagicMock,
        mock_generate_report: MagicMock,
    ) -> None:
        from src.tools.builtin import _generate_report_handler

        fake_profile = types.SimpleNamespace(enabled=True, api_base="http://example.com", api_key="k", model_name="m")
        fake_snapshot = types.SimpleNamespace(llm_profiles=[fake_profile])
        fake_service = MagicMock()
        fake_service.build_runtime_snapshot.return_value = fake_snapshot
        mock_build_settings_service.return_value = fake_service
        mock_generate_report.return_value = "/tmp/report.docx"

        with patch("httpx.post", side_effect=TimeoutError("timeout")):
            _generate_report_handler(
                {
                    "title": "车间通行和通行管理规定",
                    "content": "车间通行和通行管理规定",
                    "from": "u123",
                    "format": "docx",
                    "_context_text": (
                        "=== 【模板文件内容】===\n模板正文\n\n"
                        "=== 【用户要求】===\n请根据以下条目生成正式制度文本"
                    ),
                }
            )

        generated_content = mock_generate_report.call_args.args[1]
        self.assertIn("1目的", generated_content)
        self.assertIn("2适用范围", generated_content)
        self.assertIn("请根据以下条目生成正式制度文本", generated_content)

    def test_build_policy_fallback_content_expands_structured_request(self) -> None:
        from src.tools.builtin import _build_policy_fallback_content

        content = _build_policy_fallback_content(
            "车间通行和通讯管理规定",
            "一通行规范\n"
            "1 车间员工上下班必须在指定通道通行\n"
            "2 其他门均为应急通道，非应急情况禁止通行\n"
            "二通讯规范\n"
            "1 所有人员不得携带手机进入车间\n"
            "2 车间内配置了对讲机作为公司内部沟通使用\n",
        )

        self.assertIn("1目的", content)
        self.assertIn("4车间通行管理要求", content)
        self.assertIn("5车间通讯管理要求", content)
        self.assertIn("车间员工上下班必须在指定通道通行", content)
        self.assertIn("所有人员不得携带手机进入车间", content)
        self.assertIn("监督与处罚", content)


    def test_build_policy_fallback_content_merges_sub_items_into_parent_clause(self) -> None:
        from src.tools.builtin import _build_policy_fallback_content

        content = _build_policy_fallback_content(
            "车间通行管理规定",
            "一通行规范\n"
            "1 车间员工上下班必须在指定通道通行\n"
            "a 熔铸车间通过北面金属检测门和道闸刷脸通行\n"
            "b 其他车间和设备部人员通过车间七号门人行通道通行\n",
        )

        self.assertIn("车间通行管理要求", content)
        self.assertIn("熔铸车间通过北面金属检测门和道闸刷脸通行", content)
        self.assertIn("其他车间和设备部人员通过车间七号门人行通道通行", content)
        self.assertNotIn("4.3 a ", content)
        self.assertNotIn("4.4 b ", content)

    def test_build_policy_fallback_content_preserves_attachment_hint(self) -> None:
        from src.tools.builtin import _build_policy_fallback_content

        content = _build_policy_fallback_content(
            "车间通讯管理规定",
            "一通讯规范\n"
            "1 所有人员均不得携带手机进入车间\n"
            "2 各部门对讲机频道清单后附\n",
        )

        self.assertIn("各部门对讲机频道清单", content)
        self.assertIn("附则", content)
        self.assertIn("车间通讯管理要求", content)


class TestTemplatePreparation(unittest.TestCase):
    def setUp(self) -> None:
        self._created_paths: list[Path] = []

    def tearDown(self) -> None:
        for path in reversed(self._created_paths):
            if path.is_file():
                path.unlink(missing_ok=True)
            elif path.is_dir():
                shutil.rmtree(path, ignore_errors=True)

    def test_prepare_template_candidate_copies_docx_for_generation(self) -> None:
        from src.gateway.wecom_file_handler import prepare_template_candidate

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            source_path = Path(tmp.name)
        self._created_paths.append(source_path)
        source_path.write_bytes(b"fake-docx")

        template_info = prepare_template_candidate(str(source_path), "template.docx", "u123")
        template_path = Path(template_info["template_path"])
        self._created_paths.append(template_path.parent)

        self.assertTrue(template_info["template_candidate"])
        self.assertEqual(template_info["template_kind"], "docx_template")
        self.assertTrue(template_path.exists())
        self.assertEqual(template_path.suffix.lower(), ".docx")

        latest_file = template_path.parent / "latest_template.json"
        self.assertTrue(latest_file.exists())
        latest_payload = json.loads(latest_file.read_text(encoding="utf-8"))
        self.assertEqual(latest_payload["template_path"], str(template_path))

    def test_prepare_template_candidate_copies_xlsx_for_generation(self) -> None:
        from src.gateway.wecom_file_handler import prepare_template_candidate

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            source_path = Path(tmp.name)
        self._created_paths.append(source_path)
        source_path.write_bytes(b"fake-xlsx")

        template_info = prepare_template_candidate(str(source_path), "template.xlsx", "u123")
        template_path = Path(template_info["template_path"])
        self._created_paths.append(template_path.parent)

        self.assertTrue(template_info["template_candidate"])
        self.assertEqual(template_info["template_kind"], "xlsx_template")
        self.assertTrue(template_path.exists())
        self.assertEqual(template_path.suffix.lower(), ".xlsx")

    def test_prepare_template_candidate_copies_pptx_for_generation(self) -> None:
        from src.gateway.wecom_file_handler import prepare_template_candidate

        with tempfile.NamedTemporaryFile(suffix=".pptx", delete=False) as tmp:
            source_path = Path(tmp.name)
        self._created_paths.append(source_path)
        source_path.write_bytes(b"fake-pptx")

        template_info = prepare_template_candidate(str(source_path), "template.pptx", "u123")
        template_path = Path(template_info["template_path"])
        self._created_paths.append(template_path.parent)

        self.assertTrue(template_info["template_candidate"])
        self.assertEqual(template_info["template_kind"], "pptx_template")
        self.assertTrue(template_path.exists())
        self.assertEqual(template_path.suffix.lower(), ".pptx")

    def test_get_latest_template_candidate_returns_recent_template(self) -> None:
        from src.gateway.wecom_file_handler import get_latest_template_candidate

        with tempfile.TemporaryDirectory() as td:
            user_dir = Path(td) / "u123"
            user_dir.mkdir(parents=True)
            template_path = user_dir / "template.docx"
            template_path.write_bytes(b"fake-docx")
            latest_file = user_dir / "latest_template.json"
            latest_file.write_text(
                json.dumps(
                    {
                        "template_path": str(template_path),
                        "captured_at": 9999999999,
                        "template_kind": "docx_template",
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            with patch("src.gateway.wecom_file_handler._TEMPLATES_DIR", str(Path(td))):
                result = get_latest_template_candidate("u123", now=9999999999, max_age_seconds=60)

        self.assertEqual(result, str(template_path))

    def test_get_latest_template_candidate_filters_by_template_kind(self) -> None:
        from src.gateway.wecom_file_handler import get_latest_template_candidate

        with tempfile.TemporaryDirectory() as td:
            user_dir = Path(td) / "u123"
            user_dir.mkdir(parents=True)
            template_path = user_dir / "template.docx"
            template_path.write_bytes(b"fake-docx")
            latest_file = user_dir / "latest_template.json"
            latest_file.write_text(
                json.dumps(
                    {
                        "template_path": str(template_path),
                        "captured_at": 9999999999,
                        "template_kind": "docx_template",
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            with patch("src.gateway.wecom_file_handler._TEMPLATES_DIR", str(Path(td))):
                mismatch = get_latest_template_candidate(
                    "u123",
                    now=9999999999,
                    max_age_seconds=60,
                    expected_kind="xlsx_template",
                )
                match = get_latest_template_candidate(
                    "u123",
                    now=9999999999,
                    max_age_seconds=60,
                    expected_kind="docx_template",
                )

        self.assertIsNone(mismatch)
        self.assertEqual(match, str(template_path))

    def test_extract_docx_template_outline_includes_paragraphs_and_tables(self) -> None:
        from src.tools.document_tool import extract_docx_template_outline

        class FakeParagraph:
            def __init__(self, text: str, style_name: str) -> None:
                self.text = text
                self.style = types.SimpleNamespace(name=style_name)

        class FakeCell:
            def __init__(self, text: str) -> None:
                self.text = text

        class FakeRow:
            def __init__(self, cells: list[str]) -> None:
                self.cells = [FakeCell(cell) for cell in cells]

        class FakeTable:
            def __init__(self) -> None:
                self.rows = [
                    FakeRow(["部门", "负责人"]),
                    FakeRow(["生产部", "张三"]),
                ]

        class FakeDocument:
            def __init__(self, path: str) -> None:
                self.paragraphs = [
                    FakeParagraph("第一章 总则", "Heading 1"),
                    FakeParagraph("第一条 为了规范流程。", "Normal"),
                ]
                self.tables = [FakeTable()]

        fake_docx = types.SimpleNamespace(Document=FakeDocument)
        with patch.dict(sys.modules, {"docx": fake_docx}):
            outline = extract_docx_template_outline("fake.docx")

        self.assertIn("paragraphs", outline)
        self.assertIn("tables", outline)
        self.assertGreaterEqual(len(outline["paragraphs"]), 2)
        self.assertEqual(outline["paragraphs"][0]["style"], "Heading 1")
        self.assertEqual(outline["paragraphs"][0]["anchor"], "第一章 总则")
        self.assertEqual(len(outline["tables"]), 1)
        self.assertEqual(outline["tables"][0]["rows"], 2)
        self.assertEqual(outline["tables"][0]["cols"], 2)


class TestFileTypeInference(unittest.TestCase):
    def test_summarize_file_bytes_infers_docx_from_zip_signature(self) -> None:
        from src.gateway.wecom_file_handler import summarize_file_bytes

        office_bytes = b"PK\x03\x04" + b"[Content_Types].xml" + b"word/document.xml"

        with patch("src.knowledge.document_converter.guess_type", return_value="other"), \
             patch("src.knowledge.document_converter.convert_document", return_value="提取到的正文"), \
             patch("src.gateway.wecom_file_handler.prepare_template_candidate"), \
             patch("src.knowledge.gbrain_repo.GbrainKnowledgeRepository"), \
             patch("src.knowledge.collector.KnowledgeCollector") as mock_collector_cls:
            mock_collector_cls.return_value.collect_file.return_value = None
            summary = summarize_file_bytes(office_bytes, "7650386785395487008", from_user_id="u123")

        self.assertIn("提取到的正文", summary)

    def test_summarize_file_bytes_strips_null_characters_from_extracted_text(self) -> None:
        from src.gateway.wecom_file_handler import summarize_file_bytes

        office_bytes = b"PK\x03\x04" + b"[Content_Types].xml" + b"word/document.xml"

        with patch("src.knowledge.document_converter.guess_type", return_value="other"), \
             patch("src.knowledge.document_converter.convert_document", return_value="第一章\x00总则\x00正文"), \
             patch("src.gateway.wecom_file_handler.prepare_template_candidate"), \
             patch("src.knowledge.gbrain_repo.GbrainKnowledgeRepository"), \
             patch("src.knowledge.collector.KnowledgeCollector") as mock_collector_cls:
            mock_collector_cls.return_value.collect_file.return_value = None
            summary = summarize_file_bytes(office_bytes, "7650386785395487008", from_user_id="u123")

        self.assertNotIn("\x00", summary)
        self.assertIn("第一章总则正文", summary)


class TestDocumentContentSanitization(unittest.TestCase):
    def test_generate_report_strips_embedded_null_bytes_before_officecli(self) -> None:
        from src.tools.document_tool import generate_report

        with patch("src.tools.document_tool.os.path.isfile", return_value=True), \
             patch("src.tools.document_tool.os.makedirs"), \
             patch("src.tools.document_tool.os.path.exists", return_value=False), \
             patch("src.tools.document_tool.os.path.getsize", return_value=123), \
             patch("src.tools.document_tool._cli", return_value=(0, "")) as mock_cli:
            result = generate_report("测试文档", "第一章\x00总则\n\n第二段\x00内容", "docx")

        self.assertTrue(result.endswith(".docx"))
        joined_args = " ".join(" ".join(call.args[0]) for call in mock_cli.call_args_list)
        self.assertNotIn("\x00", joined_args)

    @patch("src.tools.document_tool._build_docx_from_template")
    def test_generate_report_uses_template_path_for_docx_when_available(self, mock_build_from_template: MagicMock) -> None:
        from src.tools.document_tool import generate_report

        with patch("src.tools.document_tool.os.path.isfile", return_value=True), \
             patch("src.tools.document_tool.os.makedirs"), \
             patch("src.tools.document_tool.os.path.exists", return_value=False), \
             patch("src.tools.document_tool.os.path.getsize", return_value=123), \
             patch("src.tools.document_tool.shutil.copyfile"), \
             patch("src.tools.document_tool._cli", return_value=(0, "")):
            result = generate_report("测试文档", "正文内容" * 10, "docx", template_path="/tmp/template.docx")

        self.assertTrue(result.endswith(".docx"))
        mock_build_from_template.assert_called_once()

    @patch("src.tools.document_tool._build_xlsx_from_template")
    def test_generate_report_uses_template_path_for_xlsx_when_available(self, mock_build_from_template: MagicMock) -> None:
        from src.tools.document_tool import generate_report

        def fake_isfile(path: str) -> bool:
            return path in {"/usr/local/bin/officecli", "/tmp/template.xlsx"}

        with patch("src.tools.document_tool.os.path.isfile", side_effect=fake_isfile), \
             patch("src.tools.document_tool.os.makedirs"), \
             patch("src.tools.document_tool.os.path.exists", return_value=False), \
             patch("src.tools.document_tool.os.path.getsize", return_value=123), \
             patch("src.tools.document_tool._cli", return_value=(0, "")):
            result = generate_report("测试文档", "A\tB\n1\t2", "xlsx", template_path="/tmp/template.xlsx")

        self.assertTrue(result.endswith(".xlsx"))
        mock_build_from_template.assert_called_once()

    @patch("src.tools.document_tool._build_pptx_from_template")
    def test_generate_report_uses_template_path_for_pptx_when_available(self, mock_build_from_template: MagicMock) -> None:
        from src.tools.document_tool import generate_report

        def fake_isfile(path: str) -> bool:
            return path in {"/usr/local/bin/officecli", "/tmp/template.pptx"}

        with patch("src.tools.document_tool.os.path.isfile", side_effect=fake_isfile), \
             patch("src.tools.document_tool.os.makedirs"), \
             patch("src.tools.document_tool.os.path.exists", return_value=False), \
             patch("src.tools.document_tool.os.path.getsize", return_value=123), \
             patch("src.tools.document_tool._cli", return_value=(0, "")):
            result = generate_report("测试文档", "第一页\n\n第二页", "pptx", template_path="/tmp/template.pptx")

        self.assertTrue(result.endswith(".pptx"))
        mock_build_from_template.assert_called_once()


class TestTemplateRequestSplit(unittest.TestCase):
    def test_split_template_and_request_preserves_full_multi_paragraph_request(self) -> None:
        from src.tools.builtin import _split_template_and_request

        content = (
            "【系统提示】你已经收到了用户上传文档的实际内容。\n\n"
            "=== 【模板文件内容】===\n"
            "用户发送了文件：模板.docx\n\n以下是文件内容：\n---\n第一章 总则\n第二章 管理要求\n---\n\n"
            "=== 【用户要求】===\n"
            "我需要你帮我生成一个车间通行和通行管理规定\n\n"
            "一通行规范\n1 车间员工上下班必须在指定通道通行\n\n"
            "二通讯规范\n1 所有人员不得携带手机进入车间"
        )

        template_text, request_text = _split_template_and_request(content)

        self.assertIn("第一章 总则", template_text)
        self.assertIn("我需要你帮我生成一个车间通行和通行管理规定", request_text)
        self.assertIn("一通行规范", request_text)
        self.assertIn("二通讯规范", request_text)


class TestStructuredRequirementSpec(unittest.TestCase):
    def test_document_requirements_module_exposes_structured_parser(self) -> None:
        from src.tools.document_requirements import build_requirement_spec

        spec = build_requirement_spec(
            "\u8f66\u95f4\u901a\u884c\u548c\u901a\u8baf\u7ba1\u7406\u89c4\u5b9a",
            "\u7b2c\u4e00\u7ae0 \u603b\u5219",
            "\u4e00\u901a\u884c\u89c4\u8303\n1 \u8f66\u95f4\u5458\u5de5\u4e0a\u4e0b\u73ed\u5fc5\u987b\u5728\u6307\u5b9a\u901a\u9053\u901a\u884c\n",
        )

        self.assertEqual(spec["sections"][0]["raw_title"], "\u901a\u884c\u89c4\u8303")
        self.assertEqual(spec["sections"][0]["heading"], "\u8f66\u95f4\u901a\u884c\u7ba1\u7406\u8981\u6c42")

    def test_build_requirement_spec_preserves_sections_items_and_subitems(self) -> None:
        from src.tools.builtin import _build_requirement_spec

        spec = _build_requirement_spec(
            "\u8f66\u95f4\u901a\u884c\u548c\u901a\u8baf\u7ba1\u7406\u89c4\u5b9a",
            "\u6a21\u677f\u793a\u4f8b\u6b63\u6587",
            "\u4e00\u901a\u884c\u89c4\u8303\n"
            "1 \u8f66\u95f4\u5458\u5de5\u4e0a\u4e0b\u73ed\u5fc5\u987b\u5728\u6307\u5b9a\u901a\u9053\u901a\u884c\n"
            "a \u7194\u94f8\u8f66\u95f4\u901a\u8fc7\u5317\u9762\u91d1\u5c5e\u68c0\u6d4b\u95e8\u548c\u9053\u95f8\u5237\u8138\u901a\u884c\n"
            "b \u5176\u4ed6\u8f66\u95f4\u548c\u8bbe\u5907\u90e8\u4eba\u5458\u901a\u8fc7\u8f66\u95f4\u4e03\u53f7\u95e8\u4eba\u884c\u901a\u9053\u901a\u884c\n"
            "2 \u529e\u516c\u5ba4\u4eba\u5458\u901a\u8fc7\u5927\u5385\u5237\u8138\u8ba4\u8bc1\u540e\u8fdb\u51fa\u8f66\u95f4\n"
            "\u4e8c\u901a\u8baf\u89c4\u8303\n"
            "1 \u6240\u6709\u4eba\u5458\u5747\u4e0d\u5f97\u643a\u5e26\u624b\u673a\u8fdb\u5165\u8f66\u95f4\n"
            "2 \u5404\u90e8\u95e8\u5bf9\u8bb2\u673a\u9891\u9053\u6e05\u5355\u540e\u9644\n",
        )

        self.assertEqual(spec["title"], "\u8f66\u95f4\u901a\u884c\u548c\u901a\u8baf\u7ba1\u7406\u89c4\u5b9a")
        self.assertEqual(len(spec["sections"]), 2)
        self.assertEqual(spec["sections"][0]["raw_title"], "\u901a\u884c\u89c4\u8303")
        self.assertEqual(spec["sections"][0]["heading"], "\u8f66\u95f4\u901a\u884c\u7ba1\u7406\u8981\u6c42")
        self.assertEqual(spec["sections"][0]["items"][0]["main"], "\u8f66\u95f4\u5458\u5de5\u4e0a\u4e0b\u73ed\u5fc5\u987b\u5728\u6307\u5b9a\u901a\u9053\u901a\u884c")
        self.assertEqual(
            spec["sections"][0]["items"][0]["sub_items"],
            [
                "\u7194\u94f8\u8f66\u95f4\u901a\u8fc7\u5317\u9762\u91d1\u5c5e\u68c0\u6d4b\u95e8\u548c\u9053\u95f8\u5237\u8138\u901a\u884c",
                "\u5176\u4ed6\u8f66\u95f4\u548c\u8bbe\u5907\u90e8\u4eba\u5458\u901a\u8fc7\u8f66\u95f4\u4e03\u53f7\u95e8\u4eba\u884c\u901a\u9053\u901a\u884c",
            ],
        )
        self.assertTrue(spec["has_attachment_hint"])

    def test_build_requirement_spec_uses_template_excerpt(self) -> None:
        from src.tools.builtin import _build_requirement_spec

        spec = _build_requirement_spec(
            "\u8f66\u95f4\u901a\u884c\u548c\u901a\u8baf\u7ba1\u7406\u89c4\u5b9a",
            "\u7b2c\u4e00\u7ae0 \u603b\u5219\n\u7b2c\u4e8c\u7ae0 \u9002\u7528\u8303\u56f4\n\u7b2c\u4e09\u7ae0 \u7ba1\u7406\u8981\u6c42",
            "\u4e00\u901a\u884c\u89c4\u8303\n1 \u8f66\u95f4\u5458\u5de5\u4e0a\u4e0b\u73ed\u5fc5\u987b\u5728\u6307\u5b9a\u901a\u9053\u901a\u884c\n",
        )

        self.assertIn("\u7b2c\u4e00\u7ae0 \u603b\u5219", spec["template_excerpt"])
        self.assertIn("\u7b2c\u4e8c\u7ae0 \u9002\u7528\u8303\u56f4", spec["template_excerpt"])

    def test_document_requirements_module_supports_notice_family(self) -> None:
        from src.tools.document_requirements import build_fallback_content

        content = build_fallback_content(
            "关于加强车间现场秩序管理的通知",
            "请发布一份通知，要求各部门本周内完成通道标识自查，并于周五前提交整改结果。",
        )

        self.assertIn("通知", content)
        self.assertIn("各部门", content)
        self.assertIn("整改结果", content)

    def test_document_requirements_module_supports_memo_family(self) -> None:
        from src.tools.document_requirements import build_fallback_content

        content = build_fallback_content(
            "车间现场秩序管理备忘录",
            "请整理一份备忘录，记录本周通道标识排查情况、已发现问题和下周整改安排。",
        )

        self.assertIn("备忘录", content)
        self.assertIn("排查情况", content)
        self.assertIn("整改安排", content)
        self.assertIn("后续安排", content)
        self.assertNotIn("监督与处罚", content)


class TestDocumentGenerationService(unittest.TestCase):
    @patch("src.tools.document_tool.generate_report")
    @patch("src.config.bootstrap.build_settings_service")
    def test_service_uses_configured_encoded_document_url(
        self,
        mock_build_settings_service: MagicMock,
        mock_generate_report: MagicMock,
    ) -> None:
        from src.tools.document_generation_service import generate_document

        fake_service = MagicMock()
        fake_service.build_runtime_snapshot.return_value = types.SimpleNamespace(llm_profiles=[])
        mock_build_settings_service.return_value = fake_service
        mock_generate_report.return_value = "/tmp/车间 管理规定.docx"

        with patch.dict(os.environ, {"ANT_COLONY_DOCUMENT_BASE_URL": "https://docs.example.test/base/"}):
            result = generate_document(
                {
                    "title": "车间管理规定",
                    "content": "这是可直接生成的正式制度内容。" * 10,
                    "format": "docx",
                }
            )

        self.assertEqual(
            result,
            "文档已生成，点击下载：https://docs.example.test/base/api/v1/documents/"
            "%E8%BD%A6%E9%97%B4%20%E7%AE%A1%E7%90%86%E8%A7%84%E5%AE%9A.docx",
        )

    @patch("src.tools.document_tool.generate_report")
    @patch("src.config.bootstrap.build_settings_service")
    def test_service_bot_file_metadata_uses_configured_document_url(
        self,
        mock_build_settings_service: MagicMock,
        mock_generate_report: MagicMock,
    ) -> None:
        from src.tools.document_generation_service import generate_document

        fake_service = MagicMock()
        fake_service.build_runtime_snapshot.return_value = types.SimpleNamespace(llm_profiles=[])
        mock_build_settings_service.return_value = fake_service
        mock_generate_report.return_value = "/tmp/车间 管理规定.docx"

        with patch.dict(os.environ, {"ANT_COLONY_DOCUMENT_BASE_URL": "https://docs.example.test"}):
            result = generate_document(
                {
                    "title": "车间管理规定",
                    "content": "这是可直接生成的正式制度内容。" * 10,
                    "format": "docx",
                    "_source_provider": "wecom_bot",
                }
            )

        self.assertTrue(result.startswith("[BOT_FILE]"))
        metadata = json.loads(result[len("[BOT_FILE]"):])
        self.assertEqual(
            metadata["download_url"],
            "https://docs.example.test/api/v1/documents/"
            "%E8%BD%A6%E9%97%B4%20%E7%AE%A1%E7%90%86%E8%A7%84%E5%AE%9A.docx",
        )

    @patch("src.tools.document_tool.generate_report")
    @patch("src.config.bootstrap.build_settings_service")
    @patch("src.gateway.wecom_outbound.send_file")
    @patch("src.gateway.wecom_outbound.send_file_card")
    def test_service_generate_document_falls_back_to_notice_content_when_enrichment_times_out(
        self,
        mock_send_file_card: MagicMock,
        mock_send_file: MagicMock,
        mock_build_settings_service: MagicMock,
        mock_generate_report: MagicMock,
    ) -> None:
        from src.tools.document_generation_service import generate_document

        fake_profile = types.SimpleNamespace(enabled=True, api_base="http://example.com", api_key="k", model_name="m")
        fake_snapshot = types.SimpleNamespace(llm_profiles=[fake_profile])
        fake_service = MagicMock()
        fake_service.build_runtime_snapshot.return_value = fake_snapshot
        mock_build_settings_service.return_value = fake_service
        mock_generate_report.return_value = "/tmp/notice.docx"
        mock_send_file_card.return_value = True

        with patch("httpx.post", side_effect=TimeoutError("timeout")):
            result = generate_document(
                {
                    "title": "关于加强车间现场秩序管理的通知",
                    "content": "关于加强车间现场秩序管理的通知",
                    "from": "u123",
                    "format": "docx",
                    "_context_text": (
                        "=== 【模板文件内容】===\n通知模板\n\n"
                        "=== 【用户要求】===\n请发布一份通知，要求各部门本周内完成通道标识自查，并于周五前提交整改结果。"
                    ),
                }
            )

        self.assertEqual(result, "")
        generated_content = mock_generate_report.call_args.args[1]
        self.assertIn("通知", generated_content)
        self.assertIn("各部门", generated_content)
        self.assertIn("整改结果", generated_content)
        mock_send_file.assert_called_once_with("u123", "/tmp/notice.docx")


class TestTemplateStructuredRendering(unittest.TestCase):
    def test_parse_docx_blocks_recognizes_markdown_tables(self) -> None:
        from src.tools.document_tool import _parse_docx_blocks

        content = (
            "## 文档发行审批\n"
            "| 角色 | 姓名 | 签字 |\n"
            "|------|------|------|\n"
            "| 起草 | 张三 |      |\n"
            "\n"
            "## 第一章 总则\n"
            "为了规范流程。"
        )

        blocks = _parse_docx_blocks(content)

        self.assertEqual(blocks[0]["type"], "heading")
        self.assertEqual(blocks[1]["type"], "table")
        self.assertEqual(blocks[1]["rows"][0], ["角色", "姓名", "签字"])
        self.assertEqual(blocks[1]["rows"][1], ["起草", "张三", ""])
        self.assertEqual(blocks[2]["type"], "heading")

    def test_preserve_table_heading_identifies_template_admin_tables(self) -> None:
        from src.tools.document_tool import _should_preserve_template_table

        self.assertTrue(_should_preserve_template_table("文档发行审批"))
        self.assertTrue(_should_preserve_template_table("文件修订履历"))
        self.assertFalse(_should_preserve_template_table("第一章 总则"))

    def test_build_docx_from_template_preserves_admin_table_and_run_formatting(self) -> None:
        from docx import Document
        from src.tools.document_tool import _build_docx_from_template

        with tempfile.TemporaryDirectory() as td:
            template_path = Path(td) / "template.docx"
            output_path = Path(td) / "output.docx"

            doc = Document()
            h = doc.add_paragraph("文档发行审批")
            h.style = "Heading 1"
            table = doc.add_table(rows=4, cols=4)
            table.cell(0, 0).text = "角色"
            table.cell(0, 1).text = "姓名"
            table.cell(1, 0).text = "起草"
            p = doc.add_paragraph("第一章 总则")
            p.style = "Heading 1"
            body = doc.add_paragraph("模板正文")
            body.style = "Normal"
            body.runs[0].bold = True
            doc.save(template_path)

            content = (
                "## 文档发行审批\n"
                "| 角色 | 姓名 | 签字 | 日期 |\n"
                "|------|------|------|------|\n"
                "| 起草 | 张三 |      |      |\n"
                "| 审核 | 李四 |      |      |\n"
                "\n"
                "## 第一章 总则\n"
                "为了规范流程。"
            )

            _build_docx_from_template(str(template_path), str(output_path), content)

            out = Document(output_path)
            self.assertEqual(len(out.tables), 1)
            self.assertEqual(out.tables[0].cell(1, 1).text, "张三")
            paragraph_texts = [p.text for p in out.paragraphs if p.text.strip()]
            self.assertIn("第一章 总则", paragraph_texts)
            body_para = next(p for p in out.paragraphs if p.text.strip() == "为了规范流程。")
            self.assertTrue(body_para.runs[0].bold)

    def test_build_docx_from_template_keeps_front_matter_before_business_sections(self) -> None:
        from copy import deepcopy
        from docx import Document
        from src.tools.document_tool import _build_docx_from_template

        with tempfile.TemporaryDirectory() as td:
            template_path = Path(td) / "template.docx"
            output_path = Path(td) / "output.docx"

            doc = Document()
            doc.add_paragraph("企业内部沟通管理办法")
            front_heading = doc.add_paragraph("文档发行审批")
            front_heading.style = "Heading 1"
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "角色"
            table.cell(0, 1).text = "姓名"
            table.cell(1, 0).text = "起草"
            table.cell(1, 1).text = "张三"
            business_heading = doc.add_paragraph("1目的")
            business_heading.style = "Heading 1"
            doc.add_paragraph("旧正文内容")
            doc.save(template_path)

            content = (
                "## 第一章 总则\n"
                "新的正文内容。\n\n"
                "## 第二章 通讯规范\n"
                "新的第二章内容。"
            )

            _build_docx_from_template(str(template_path), str(output_path), content)

            out = Document(output_path)
            texts = [p.text.strip() for p in out.paragraphs if p.text.strip()]
            self.assertEqual(texts[0], "企业内部沟通管理办法")
            self.assertEqual(texts[1], "文档发行审批")
            self.assertIn("第一章 总则", texts)
            self.assertNotIn("旧正文内容", texts)
            self.assertEqual(len(out.tables), 1)
            self.assertEqual(out.tables[0].cell(1, 1).text, "张三")



    def test_build_docx_from_template_preserves_paragraph_level_formatting(self) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
        from src.tools.document_tool import _build_docx_from_template

        with tempfile.TemporaryDirectory() as td:
            template_path = Path(td) / "template.docx"
            output_path = Path(td) / "output.docx"

            doc = Document()
            heading = doc.add_paragraph("1 Purpose")
            heading.style = "Heading 1"
            body = doc.add_paragraph("Old body content")
            body.style = "Normal"
            body.alignment = WD_ALIGN_PARAGRAPH.CENTER
            body.paragraph_format.left_indent = Inches(0.5)
            body.paragraph_format.space_before = Pt(18)
            body.paragraph_format.space_after = Pt(12)
            doc.save(template_path)

            content = (
                "## First Section\n"
                "New body content."
            )

            _build_docx_from_template(str(template_path), str(output_path), content)

            out = Document(output_path)
            body_para = next(p for p in out.paragraphs if p.text.strip() == "New body content.")
            self.assertEqual(body_para.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertIsNotNone(body_para.paragraph_format.left_indent)
            self.assertAlmostEqual(body_para.paragraph_format.left_indent.inches, 0.5, places=2)
            self.assertEqual(body_para.paragraph_format.space_before.pt, 18)
            self.assertEqual(body_para.paragraph_format.space_after.pt, 12)


class TestTemplatePromptConstruction(unittest.TestCase):
    def test_build_template_prompt_block_includes_template_excerpt(self) -> None:
        from src.tools.builtin import _build_template_prompt_block

        with patch(
            "src.tools.document_tool.extract_docx_template_outline",
            return_value={
                "paragraphs": [{"text": "First Section", "style": "Heading 1"}],
                "tables": [],
            },
        ):
            block = _build_template_prompt_block(
                "/tmp/template.docx",
                "Template body excerpt\nClause one content",
            )

        self.assertIn("Heading 1", block)
        self.assertIn("Template body excerpt", block)
        self.assertIn("Clause one content", block)


class TestWeComFileRouting(unittest.TestCase):
    def test_build_gateway_payload_uses_file_name_for_file_messages(self) -> None:
        from src.gateway.wecom_callback_server import build_gateway_payload

        payload = build_gateway_payload(
            {
                "MsgType": "file",
                "FromUserName": "u123",
                "MsgId": "mid-1",
                "MediaId": "media-1",
                "FileName": "模板.docx",
            }
        )

        self.assertEqual(payload["msg_type"], "file")
        self.assertEqual(payload["content"], "模板.docx")
        self.assertEqual(payload["media_id"], "media-1")

    def test_build_gateway_payload_falls_back_to_file_name_when_content_empty(self) -> None:
        from src.gateway.wecom_callback_server import build_gateway_payload

        payload = build_gateway_payload(
            {
                "MsgType": "file",
                "FromUserName": "u123",
                "MsgId": "mid-2",
                "MediaId": "media-2",
                "Content": "",
                "FileName": "制度.docx",
            }
        )

        self.assertEqual(payload["content"], "制度.docx")

    def test_file_message_after_recent_text_uses_buffered_text(self) -> None:
        from src.gateway.inbound_service import InboundGatewayService, _text_buffer, _file_buffer
        from src.gateway.dispatcher import Dispatcher
        from src.models.contracts import AgentResponse

        _text_buffer.clear()
        _file_buffer.clear()
        self.addCleanup(_text_buffer.clear)
        self.addCleanup(_file_buffer.clear)

        service = InboundGatewayService(dispatcher=Dispatcher(), batch_processor=MagicMock())
        service._conversations = MagicMock()
        service._conversations.get.return_value = MagicMock(get_context=MagicMock(return_value=""), add=MagicMock())
        service._conversations.save_all = MagicMock()
        fake_agent = MagicMock()
        fake_agent.process_message.return_value = AgentResponse(text="combined ok")
        service.get_or_create_agent = MagicMock(return_value=fake_agent)

        _text_buffer["u123"] = ("请提炼这个模板的重点", 0.0)

        with patch("src.gateway.inbound_service._time.time", return_value=1.0), \
             patch("src.gateway.wecom_file_handler.handle_wecom_file", return_value={"summary": "用户发送了文件：模板.docx\n\n以下是文件内容：\n---\n第一章 总则\n---\n", "template_path": "/tmp/template.docx"}):
            result = service.handle_wecom_payload(
                {
                    "from_user_id": "u123",
                    "msg_type": "file",
                    "media_id": "media-1",
                    "content": "模板.docx",
                    "is_direct": True,
                }
            )

        self.assertEqual(result.route_kind, "personal")
        self.assertEqual(result.response.text, "combined ok")
        combined_text = fake_agent.process_message.call_args.args[1]
        self.assertIn("用户发送了文件：模板.docx", combined_text)
        self.assertIn("请提炼这个模板的重点", combined_text)
        self.assertNotIn("u123", _text_buffer)

    def test_file_message_ignores_stale_conversation_context(self) -> None:
        from src.gateway.inbound_service import InboundGatewayService, _text_buffer, _file_buffer
        from src.gateway.dispatcher import Dispatcher
        from src.models.contracts import AgentResponse

        _text_buffer.clear()
        _file_buffer.clear()
        self.addCleanup(_text_buffer.clear)
        self.addCleanup(_file_buffer.clear)

        service = InboundGatewayService(dispatcher=Dispatcher(), batch_processor=MagicMock())
        fake_convo = MagicMock()
        fake_convo.get_context.return_value = "上一份文件：车间通行管理规定"
        fake_convo.add = MagicMock()
        service._conversations = MagicMock()
        service._conversations.get.return_value = fake_convo
        service._conversations.save_all = MagicMock()
        fake_agent = MagicMock()
        fake_agent.process_message.return_value = AgentResponse(text="ok")
        service.get_or_create_agent = MagicMock(return_value=fake_agent)
        _file_buffer["u123"] = ({"summary": "用户发送了文件：制度.docx\n\n以下是文件内容：\n---\n第一章 总则\n---"}, 1.0)

        with patch("src.gateway.inbound_service._time.time", return_value=2.0):
            service.handle_wecom_payload(
                {
                    "from_user_id": "u123",
                    "msg_type": "text",
                    "content": "分析这个文档的内容",
                    "is_direct": True,
                }
            )

        self.assertEqual(fake_agent.process_message.call_args.kwargs["conversation_context"], "")

    def test_file_message_with_generation_intent_bypasses_llm_and_generates_document(self) -> None:
        from src.gateway.inbound_service import InboundGatewayService, _text_buffer, _file_buffer
        from src.gateway.dispatcher import Dispatcher

        _text_buffer.clear()
        _file_buffer.clear()
        self.addCleanup(_text_buffer.clear)
        self.addCleanup(_file_buffer.clear)

        service = InboundGatewayService(dispatcher=Dispatcher(), batch_processor=MagicMock())
        service._conversations = MagicMock()
        service._conversations.get.return_value = MagicMock(get_context=MagicMock(return_value=""), add=MagicMock())
        service._conversations.save_all = MagicMock()
        service.get_or_create_agent = MagicMock()
        _file_buffer["u123"] = ({"summary": "用户发送了文件：模板.docx\n\n以下是文件内容：\n---\n第一章 总则\n---", "template_path": "/tmp/template.docx"}, 1.0)

        with patch("src.gateway.inbound_service._time.time", return_value=2.0), \
             patch("src.tools.builtin._generate_report_handler", return_value="[BOT_FILE]{\"path\":\"/tmp/report.docx\"}") as mock_generate:
            result = service.handle_wecom_payload(
                {
                    "from_user_id": "u123",
                    "msg_type": "text",
                    "content": "我需要你帮我生成一个车间通行和通讯管理规定",
                    "is_direct": True,
                    "provider": "wecom_bot",
                }
            )

        self.assertEqual(result.route_kind, "personal")
        self.assertEqual(result.response.text, "[BOT_FILE]{\"path\":\"/tmp/report.docx\"}")
        self.assertFalse(service.get_or_create_agent.called)
        args = mock_generate.call_args.args[0]
        self.assertEqual(args["from"], "u123")
        self.assertEqual(args["_source_provider"], "wecom_bot")
        self.assertEqual(args["_template_path"], "/tmp/template.docx")
        self.assertIn("用户发送了文件：模板.docx", args["_context_text"])
